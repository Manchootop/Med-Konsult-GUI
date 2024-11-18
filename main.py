from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.label import Label
from kivy.uix.scrollview import ScrollView
from kivy.uix.gridlayout import GridLayout
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup

import logging
import openpyxl
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')


class ExcelDocxProcessorApp(App):
    def build(self):
        layout = BoxLayout(orientation='vertical')

        # File upload button
        self.upload_button = Button(text="Upload Excel File", size_hint=(1, None), height=40)
        self.upload_button.bind(on_press=self.open_filechooser)
        layout.add_widget(self.upload_button)

        # Button to process the Excel file
        process_excel_btn = Button(text='Process Excel', size_hint=(1, None), height=40)
        process_excel_btn.bind(on_press=self.process_excel)
        layout.add_widget(process_excel_btn)

        # Input fields for docx filling (assuming two templates)
        self.template_1_input = TextInput(hint_text='Template 1 info', size_hint=(1, None), height=40, multiline=False, )
        self.template_2_input = TextInput(hint_text='Template 2 info', size_hint=(1, None), height=40, multiline=False)
        layout.add_widget(self.template_1_input)
        layout.add_widget(self.template_2_input)

        # Button to fill the docx files
        process_docx_btn = Button(text='Fill Docx', size_hint=(1, None), height=40)
        process_docx_btn.bind(on_press=self.fill_docx)
        layout.add_widget(process_docx_btn)

        # Log area
        self.log_area = ScrollView(size_hint=(1, None), height=200)
        self.log_box = GridLayout(cols=1, size_hint_y=None)
        self.log_box.bind(minimum_height=self.log_box.setter('height'))
        self.log_area.add_widget(self.log_box)
        layout.add_widget(self.log_area)

        # Bind tab navigation to the Window
        self.tab_order = [self.template_1_input, self.template_2_input]
        Window.bind(on_key_down=self.on_key_down)

        return layout

    def on_key_down(self, window, key, scancode, codepoint, modifier):
        if key == 9:  # Tab key
            focused_widget = next((widget for widget in self.tab_order if widget.focus), None)
            if focused_widget:
                current_index = self.tab_order.index(focused_widget)
                next_index = (current_index + 1) % len(self.tab_order)
                self.tab_order[next_index].focus = True
                return True
        return False

    def open_filechooser(self, instance):
        # Create a FileChooser popup
        content = BoxLayout(orientation='vertical')
        filechooser = FileChooserIconView()
        content.add_widget(filechooser)

        # Create buttons for confirming selection
        button_layout = BoxLayout(size_hint_y=None, height=50)
        button_layout.add_widget(Button(text="Cancel", on_press=self.close_popup))
        select_button = Button(text="Select", on_press=lambda x: self.select_file(filechooser.selection))
        button_layout.add_widget(select_button)
        content.add_widget(button_layout)

        # Create and show the popup
        self.filechooser_popup = Popup(title="Select Excel File", content=content, size_hint=(0.8, 0.8))
        self.filechooser_popup.open()

    def close_popup(self, instance):
        # Close the file chooser popup
        self.filechooser_popup.dismiss()

    def select_file(self, selection):
        if selection:
            # If a file is selected, set the path to the file input
            self.excel_file_path = selection[0]
            logging.info(f"Selected Excel file: {self.excel_file_path}")
            self.log(f"Selected Excel file: {self.excel_file_path}")
        else:
            logging.warning("No file selected")
            self.log("No file selected")
        # Close the file chooser popup
        self.close_popup(instance=None)

    def process_excel(self, instance):
        # Log input action
        if hasattr(self, 'excel_file_path') and self.excel_file_path:
            logging.info(f"Processing Excel file: {self.excel_file_path}")
            self.log(f"Processing Excel file: {self.excel_file_path}")
            # Read Excel file
            wb = openpyxl.load_workbook(self.excel_file_path)
            sheet = wb.active
            # Process data (for simplicity, just read the first column)
            for row in sheet.iter_rows(min_row=2, min_col=1, max_col=1):
                for cell in row:
                    logging.info(f"Read data from Excel: {cell.value}")
                    self.log(f"Read data from Excel: {cell.value}")
        else:
            logging.warning("No Excel file selected")
            self.log("No Excel file selected")

    def fill_docx(self, instance):
        # Log filling action
        template_1_info = self.template_1_input.text
        template_2_info = self.template_2_input.text
        if template_1_info and template_2_info:
            logging.info(f"Filling template 1 with: {template_1_info}")
            logging.info(f"Filling template 2 with: {template_2_info}")
            self.log(f"Filling template 1 with: {template_1_info}")
            self.log(f"Filling template 2 with: {template_2_info}")
            # Process Docx templates
            self.create_docx(template_1_info, template_2_info)
        else:
            logging.warning("Incomplete template information")
            self.log("Incomplete template information")

    def create_docx(self, template_1_info, template_2_info):
        # Create first document from template
        doc1 = Document()
        doc1.add_paragraph("Template 1 filled with: " + template_1_info)
        doc1.save("template_1_filled.docx")
        logging.info("Saved Template 1 as 'template_1_filled.docx'")
        self.log("Saved Template 1 as 'template_1_filled.docx'")

        # Create second document from template
        doc2 = Document()
        doc2.add_paragraph("Template 2 filled with: " + template_2_info)
        doc2.save("template_2_filled.docx")
        logging.info("Saved Template 2 as 'template_2_filled.docx'")
        self.log("Saved Template 2 as 'template_2_filled.docx'")

    def log(self, message):
        # Add log message to the log area
        self.log_box.add_widget(Label(text=message, size_hint_y=None, height=40))


if __name__ == '__main__':
    ExcelDocxProcessorApp().run()
