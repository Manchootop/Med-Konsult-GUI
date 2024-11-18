from docx import Document

def save_screening_as_docx(id_number, data):
    # Ensure id_number is a string to slice safely
    id_number_str = str(id_number)
    last_4_digits = id_number_str[-4:]  # Extract last 4 digits

    # File name
    file_name = f"скрининг_{last_4_digits}.docx"

    # Create a Word document
    doc = Document()

    # Add a heading (optional)
    doc.add_heading("Скрининг Документ", level=1)

    # Add a table
    table = doc.add_table(rows=1, cols=len(data[0]))  # Use the number of columns in the data

    # Add table headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(data[0]):
        header_cells[i].text = header

    # Add the rest of the rows
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)  # Ensure all data is string for Word compatibility

    # Save the document
    doc.save(file_name)
    print(f"Document saved as {file_name}")

# Example usage
id_number = 1234567890
data = [
    ["Name", "Age", "Score"],
    ["Alice", 30, 85],
    ["Bob", 25, 90],
    ["Charlie", 35, 88],
]

save_screening_as_docx(id_number, data)
