from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.scrollview import ScrollView
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
from kivy.core.window import Window
from kivy.utils import get_color_from_hex
from kivy.graphics import Color, RoundedRectangle
from docx import Document
import os


class StyledTextInput(TextInput):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.size_hint = (1, None)
        self.height = 50
        self.font_size = 16
        self.background_color = get_color_from_hex("#F5F5F5")
        self.foreground_color = get_color_from_hex("#000000")
        self.hint_text_color = get_color_from_hex("#A9A9A9")
        self.padding = [10, 15, 10, 15]


class StyledButton(Button):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.size_hint = (1, None)
        self.height = 50
        self.font_size = 16
        self.background_color = get_color_from_hex("#4CAF50")
        self.color = get_color_from_hex("#FFFFFF")
        self.bold = True


class StyledLabel(Label):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.font_size = 16
        self.color = get_color_from_hex("#000000")
        self.padding = [10, 10]


class ScreeningFormApp(App):

    def build(self):
        # Set the background color
        Window.clearcolor = get_color_from_hex("#E8F5E9")

        # Main layout
        layout = BoxLayout(orientation='vertical', padding=20, spacing=10)

        # Header
        header = StyledLabel(text="Screening Form", size_hint=(1, None), height=60, bold=True, font_size=24)
        header.color = get_color_from_hex("#2E7D32")
        layout.add_widget(header)

        # Creating input fields with styles
        self.names_input = StyledTextInput(hint_text="Enter Patient's Full Name")
        self.dob_input = StyledTextInput(hint_text="Enter Date of Birth (DD.MM.YYYY)")
        self.doctor_input = StyledTextInput(hint_text="Enter Doctor's Name")
        self.coordinator_input = StyledTextInput(hint_text="Enter Coordinator's Name")
        self.screening_time_input = StyledTextInput(hint_text="Enter Screening Time (HH:MM)")
        self.age_input = StyledTextInput(hint_text="Enter Patient's Age")
        self.id_input = StyledTextInput(hint_text="Enter Patient's ID Number")

        # List of input fields for tab navigation
        self.input_fields = [
            self.names_input,
            self.dob_input,
            self.doctor_input,
            self.coordinator_input,
            self.screening_time_input,
            self.age_input,
            self.id_input,
        ]

        # Adding input fields
        for field in self.input_fields:
            layout.add_widget(field)

        # Buttons
        submit_button = StyledButton(text="Generate Document", on_press=self.generate_document)
        preview_button = StyledButton(text="Preview Word Document", on_press=self.open_file_chooser)

        layout.add_widget(submit_button)
        layout.add_widget(preview_button)

        # Scroll view for displaying results
        scroll_view = ScrollView(size_hint=(1, None), size=(Window.width, 200))
        with scroll_view.canvas.before:
            Color(*get_color_from_hex("#FFFFFF"))
            self.rect = RoundedRectangle(size=scroll_view.size, pos=scroll_view.pos, radius=[10])

        self.result_label = StyledLabel(size_hint_y=None, height=200, text='Generated Document will appear here...', valign="top")
        scroll_view.add_widget(self.result_label)

        layout.add_widget(scroll_view)

        return layout

    def on_start(self):
        # Bind Tab navigation
        Window.bind(on_key_down=self.on_key_down)

    def on_key_down(self, window, key, scancode, codepoint, modifiers):
        if key == 9:  # Tab key
            focused_widget = next((widget for widget in self.input_fields if widget.focus), None)
            if focused_widget:
                current_index = self.input_fields.index(focused_widget)
                next_index = (current_index + 1) % len(self.input_fields)
                self.input_fields[next_index].focus = True
                return True
        return False

    def generate_document(self, instance):
        # Retrieve input values
        names = self.names_input.text
        dob = self.dob_input.text
        doctor = self.doctor_input.text
        coordinator = self.coordinator_input.text
        screening_time = self.screening_time_input.text
        age = self.age_input.text
        patient_id = self.id_input.text

        # Updated Document Template
        document_template = f"""
    СПОНСОР: ModernaTX, Inc.
    ЦЕНТЪР: BG004                                                                                              ДАТА: 12.11.2024
    ПРОТОКОЛ №: mRNA-1010-P304                               Пациент: {names}
    ГЛАВЕН ИЗСЛЕДОВАТЕЛ: Проф. Д-р М. Цекова            Номер пациент: {patient_id}

    СКРИНИНГ ВИЗИТА 
    На 12.11.2024 год. в {screening_time}ч. пациентът - {names} д.р. {dob} год. се яви в „Медицински център Медконсулт Плевен“. Пациентът е съгласен да участва в клиничното изпитване след като се е запознал подробно вкъщи с дадения му формуляр ИС на хартиен носител - основен ФИС на български език версия 1.1 от 06 септември 2024 г. въз основа на 285485 mRNA-1010-P304 Основен ФИС версия 3.0 от 14 март 2024г. Пациентът бе информиран, че участието му в клиничното изпитване е доброволно, като му бяха разяснени всички ползи и рискове, както и възможността на пациента да оттегли съгласието си по всяко време, както и възможността за лечение ако не бъде включен в изпитването.
    Проведе се дискусия и бе отговорено на всичките му въпроси, след което пациентът подписа и датира два екземпляра на ИС в {screening_time}ч. на 12.11.2024 год. ИС беше подписано и датирано и от {doctor}. Единият екземпляр беше предоставен на пациента. Другият екземпляр остава в папката на пациента по клиничното проучване.
    Пациентът не желае личният му лекар да бъде информиран за участието му в клиничното изпитване.
    Демографски данни: Дата на раждане: {dob} год., възраст – {age} год. Пол – мъж, Раса: Бяла, Етнос – неиспански или латино.
    Медицинска и Ваксинационна История: (по данни на пациента).
    АХ от 2014  г.
    Пациентът съобщава, че му не му е поставяна противогрипна лицензирана или клинично изпитвана ваксина от м. август 2023 год. до момента.
    Пациентът съобщава, че не е приемал кортикостероиди и не са му прилагани имуноглобулини или кръвни продукти в последните 90 дни.
    Пациентът съобщава, че не е приемал системна имунносупресивна или имунномоделираща терапия, дългодействаща биологична терапия с имунна реакция в последните 180 дни.
    Медицински интервенции/операции –  няма.
    Минала и съпътстваща терапия:
    Лориста 50 мг 1т.дн.по повод АХ от 2018 г.
    Пушач – пуши по 10 цигари дневно от 2009 г. до сега.
    Пациентът не употребява и не е употребявал алкохол и наркотични вещества
    .... ч. Измериха  се: ръст 172,0  см., тегло 87,4 кг.
    Физикален преглед: Мъж на възраст, отговаряща на действителната, в добро общо състояние, ориентирана за време, място и собствена личност,  заема активно положение в леглото. Кожа и видими лигавици – бледорозови. Умерено изразена подкожна мастна тъкан. Език – влажен, необложен. Без шиен венозен застой. Щит. жлеза – при палпация -  неувеличена. ПЛВ – не се палпират увеличени. ДС –симетричен гръден кош, двустранно  везикуларно дишане без хрипове. ССС- Аритмична неучестена сърдечна дейност, ясни тонове. Корем на нивото на гръдния кош, меки коремни стени без палпаторна болезненост, черен дроб и слезка – не се палпират увеличени. Крайници – без отоци, запазени пулсации на периферни съдове. Нормален неврологичен статус.
    ...... ч. Измерени АН и пулс в седнало положение след 5 мин. в покой:  АН 138 /84 mmHg на доминантна дясна ръка, пулс 68 уд/мин. Дихателна честота 19 в мин, телесна температура –...
    """

        # Display the result
        self.result_label.text = document_template

        # Save the document
        self.save_document(patient_id, document_template)

    def save_document(self, patient_id, document_text):
        doc = Document()
        doc.add_paragraph(document_text)
        patient_id_str = str(patient_id)  # Ensure it's a string
        last_4_digits = patient_id_str[-4:]

        file_name = f"скрининг_{last_4_digits}.docx"
        file_path = os.path.join(os.getcwd(), file_name)
        doc.save(file_path)
        print(f"Document saved as {file_path}")

    def open_file_chooser(self, instance):
        file_chooser = FileChooserIconView(filters=['*.docx'])
        popup = Popup(title="Select Word Document", content=file_chooser, size_hint=(0.8, 0.8))
        open_button = StyledButton(text="Open", on_press=lambda x: self.load_word_file(file_chooser.selection, popup))
        popup.content.add_widget(open_button)
        popup.open()

    def load_word_file(self, selected, popup):
        if selected:
            word_file = selected[0]
            document_text = self.extract_text_from_word(word_file)
            self.result_label.text = document_text
            popup.dismiss()

    def extract_text_from_word(self, word_file):
        document = Document(word_file)
        return '\n'.join(para.text for para in document.paragraphs)


if __name__ == "__main__":
    ScreeningFormApp().run()
